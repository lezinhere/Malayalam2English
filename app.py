import streamlit as st
import torch
from transformers import AutoModelForSeq2SeqLM, AutoTokenizer
from IndicTransToolkit import IndicProcessor
from docx import Document
from io import BytesIO
import time

# ===== CONFIG =====
MODEL_NAME = "ai4bharat/indictrans2-indic-en-dist-200M"
SRC_LANG = "mal_Mlym"
TGT_LANG = "eng_Latn"
BATCH_SIZE = 4
MAX_FILE_SIZE_MB = 10

# ===== PAGE SETTINGS =====
st.set_page_config(
    page_title="Malayalam → English Translator",
    layout="centered"
)

st.title("📘 Malayalam → English Book Translator")
st.markdown("Upload a Malayalam `.docx` file to translate it into English.")

# ===== LOAD MODEL (Cached to avoid reloading) =====
@st.cache_resource
def load_model():
    tokenizer = AutoTokenizer.from_pretrained(MODEL_NAME, trust_remote_code=True)
    model = AutoModelForSeq2SeqLM.from_pretrained(MODEL_NAME, trust_remote_code=True)
    ip = IndicProcessor(inference=True)
    model.to("cpu")
    return tokenizer, model, ip

tokenizer, model, ip = load_model()

# ===== FILE UPLOAD =====
uploaded_file = st.file_uploader("Upload DOCX file", type="docx")

if uploaded_file:

    # ===== FILE SIZE VALIDATION =====
    file_size_mb = uploaded_file.size / (1024 * 1024)

    if file_size_mb > MAX_FILE_SIZE_MB:
        st.error(f"❌ File too large! Maximum allowed size is {MAX_FILE_SIZE_MB} MB.")
        st.stop()

    st.info("📖 Reading document...")

    doc = Document(uploaded_file)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    total_paragraphs = len(paragraphs)
    total_words = sum(len(p.split()) for p in paragraphs)

    st.write(f"📄 Paragraphs detected: {total_paragraphs}")
    st.write(f"📝 Word Count: {total_words}")

    # ===== ESTIMATED TIME =====
    estimated_seconds = total_paragraphs * 1.5
    estimated_minutes = round(estimated_seconds / 60, 1)

    st.info(f"⏳ Estimated translation time: ~ {estimated_minutes} minutes")

    # ===== START BUTTON =====
    if st.button("🚀 Start Translation"):

        progress_bar = st.progress(0)
        new_doc = Document()

        st.info("🔄 Translating... Please wait.")
        start_time = time.time()

        for i in range(0, total_paragraphs, BATCH_SIZE):

            batch = paragraphs[i:i+BATCH_SIZE]

            batch_processed = ip.preprocess_batch(
                batch,
                src_lang=SRC_LANG,
                tgt_lang=TGT_LANG,
            )

            inputs = tokenizer(
                batch_processed,
                truncation=True,
                padding="longest",
                return_tensors="pt",
            ).to("cpu")

            with torch.no_grad():
                generated_tokens = model.generate(
                    **inputs,
                    max_length=256,
                    num_beams=3
                )

            decoded = tokenizer.batch_decode(
                generated_tokens.detach().cpu(),
                skip_special_tokens=True
            )

            translations = ip.postprocess_batch(decoded, lang=TGT_LANG)

            for t in translations:
                new_doc.add_paragraph(t)

            progress_bar.progress(min((i + BATCH_SIZE) / total_paragraphs, 1.0))

        end_time = time.time()
        actual_minutes = round((end_time - start_time) / 60, 2)

        st.success(f"✅ Translation completed in {actual_minutes} minutes!")

        # ===== PREPARE DOWNLOAD =====
        output_buffer = BytesIO()
        new_doc.save(output_buffer)
        output_buffer.seek(0)

        st.download_button(
            label="📥 Download Translated Document",
            data=output_buffer,
            file_name="translated_english_book.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

# ===== FOOTER =====
st.markdown(
    """
    <hr style="margin-top:50px;">
    <p style="text-align:center; font-size:13px; color:#888;">
        Developed by <b>Lezin</b>
    </p>
    """,
    unsafe_allow_html=True
)
